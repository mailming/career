"""Generate Mingresume.docx from resume.md.

Reads resume.md, parses the markdown structure we use, and writes a Word
document that reuses the styles already defined in the existing
Mingresume.docx template (Title, Heading 1/2, Resume Bullets, Normal).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
SRC_MD = ROOT / "resume.md"
TEMPLATE = ROOT / "Mingresume.docx"
OUT = ROOT / "Mingresume.docx"


def md_inline_runs(paragraph, text: str) -> None:
    """Render a single line of markdown-ish text with **bold**, *italic*, `code`."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*\s][^*]*\*|(?<!\w)_[^_\s][^_]*_(?!\w))"
    )
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:  # *italic* or _italic_
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(2)
    p_format.space_after = Pt(2)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def main() -> None:
    md = SRC_MD.read_text().splitlines()

    # Reuse the existing docx as a template so we keep styles like "Resume
    # Bullets" that were defined in it. We then strip its content.
    doc = Document(str(TEMPLATE))

    # Wipe the existing body content (paragraphs + tables) without removing
    # the section properties that follow them.
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}", 1)[-1]
        if tag in {"p", "tbl"}:
            body.remove(child)

    styles = doc.styles
    style_names = {s.name for s in styles}

    def ensure_style(name: str) -> str:
        return name if name in style_names else "Normal"

    bullet_style = ensure_style("Resume Bullets")

    # Page margins: 0.7" all around for a denser, professional look.
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    state = {"in_table": False, "table_rows": [], "skip_blank_after_rule": False}

    def flush_table() -> None:
        if not state["table_rows"]:
            return
        rows = state["table_rows"]
        header = rows[0]
        data_rows = rows[2:] if len(rows) >= 2 else []
        n_cols = len(header)
        table = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
        table.autofit = False
        if "Table Grid" in style_names:
            table.style = "Table Grid"
        # Apply manual borders so the table always renders with grid lines.
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "BFBFBF")
            tblBorders.append(b)
        # Remove any existing borders before adding the new ones.
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblBorders)
        # Set column widths
        left_width = Inches(1.7)
        right_width = Inches(5.3)
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = left_width if i == 0 else right_width
        # Header
        for i, cell in enumerate(table.rows[0].cells):
            p = cell.paragraphs[0]
            run = p.add_run(header[i])
            run.bold = True
            run.font.size = Pt(10)
        # Body
        for r_idx, r in enumerate(data_rows, start=1):
            for i, cell_text in enumerate(r):
                if i >= n_cols:
                    break
                cell = table.rows[r_idx].cells[i]
                p = cell.paragraphs[0]
                md_inline_runs(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(10)
        state["in_table"] = False
        state["table_rows"] = []

    i = 0
    while i < len(md):
        line = md[i].rstrip()

        # Detect markdown table
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            state["in_table"] = True
            state["table_rows"].append(cells)
            i += 1
            continue
        elif state["in_table"]:
            flush_table()

        if not line:
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph(line[2:].strip())
            p.style = doc.styles[ensure_style("Title")]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(22)
                run.bold = True
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_paragraph(line[3:].strip().upper())
            p.style = doc.styles[ensure_style("Heading 1")]
            for run in p.runs:
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if line.startswith("### "):
            heading_style = ensure_style("Resume Header")
            p = doc.add_paragraph()
            p.style = doc.styles[heading_style]
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(0)
            i += 1
            continue

        if line.strip() == "---":
            add_horizontal_rule(doc)
            i += 1
            continue

        if line.lstrip().startswith("- "):
            content = line.lstrip()[2:]
            p = doc.add_paragraph(style=bullet_style)
            md_inline_runs(p, content)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(10.5)
            i += 1
            continue

        # Plain paragraph (handles things like the **Insurance · 06/2025 - Present** line)
        text = line
        # Markdown trailing two-space line break is unimportant for docx.
        p = doc.add_paragraph()
        md_inline_runs(p, text)
        # Center the contact line block at the top (lines 2 of original markdown)
        if (
            "@" in text
            and "·" in text
            and i < 10
        ):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("**") and "·" in text:
            # Job-meta line; italicize-ish via smaller size
            for run in p.runs:
                run.font.size = Pt(10)
                run.italic = True
        else:
            for run in p.runs:
                if not run.font.size:
                    run.font.size = Pt(10.5)
        p.paragraph_format.space_after = Pt(2)
        i += 1

    if state["in_table"]:
        flush_table()

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
